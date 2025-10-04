# ocr_any.py
import os
import sys
import io
import re
import json
import base64
import mimetypes
from pathlib import Path
from typing import Dict, List

from openai import OpenAI
from PIL import Image
import fitz  # PyMuPDF
from docx import Document
from docx.opc.constants import RELATIONSHIP_TYPE as RT

# ========= Config via env =========
MODEL = os.getenv("OPENAI_MODEL", "gpt-4o")
MODEL_FALLBACK = os.getenv("OPENAI_MODEL_FALLBACK", "gpt-4.1-mini")
OCR_LANG = os.getenv("OCR_LANG", "en").strip().lower()
if OCR_LANG not in {"en", "hi"}:
    OCR_LANG = "en"

# Modes
OCR_REDACT = os.getenv("OCR_REDACT", "false").strip().lower() == "true"
OCR_TRANSLATE_ALWAYS = os.getenv("OCR_TRANSLATE_ALWAYS", "false").strip().lower() == "true"
OCR_STRUCTURED = os.getenv("OCR_STRUCTURED", "false").strip().lower() == "true"

# Prompts (free-text OCR mode)
PROMPT_OCR_VERBATIM_EN = "OCR this image. Return the text exactly as printed in English."
PROMPT_OCR_HI_TO_EN = "OCR this image. Extract the Hindi (Devanagari) text and return an accurate English translation ONLY. Preserve structure and line breaks where possible."
PROMPT_OCR_ANY_TO_EN = "OCR this image. Extract any visible text (any language) and return an accurate English translation ONLY. Preserve structure and line breaks where possible."

# Redaction guidance (optional)
REDACTION_GUIDE = (
    "If you encounter government IDs or highly sensitive numbers, redact them by masking all "
    "but the last 4 characters (e.g., 'XXXXXXXXXXXX1234'). For PAN (pattern like AAAAA9999A), "
    "mask as 'XXXXX9999X'. Keep everything else verbatim."
)

# Systems
SYSTEM_BASE = (
    "You are an OCR engine. The user has provided this document and explicitly consents to OCR. "
    "Your ONLY task is to transcribe text from the image(s) or page(s). "
    "Do not add commentary, summaries, or warnings. Do not refuse unless the image is clearly illegal content. "
)
SYSTEM_VERBATIM = SYSTEM_BASE + "Return plain text only, preserving line breaks."
SYSTEM_REDACT = SYSTEM_BASE + REDACTION_GUIDE + " Return plain text only, preserving line breaks."
SYSTEM_OCR_TRANSLATE = (
    "You are an OCR engine with translation. The user has provided this document and explicitly consents to OCR. "
    "Your ONLY task is to transcribe visible text from the image(s) and output an English translation ONLY. "
    "Preserve formatting and line breaks as much as possible. Do not add commentary."
)
SYSTEM_TRANSLATOR = (
    "You are a professional translator. Translate the provided text into natural, fluent English. "
    "Preserve formatting and line breaks as much as possible. Do not add commentary."
)

# Structured extraction system
SYSTEM_STRUCTURED = (
    "You are an information extraction engine. Extract ONLY the requested fields from the provided content. "
    "Return a strict JSON object containing ONLY the keys from the provided schema. "
    "Omit any field that is not confidently present. Do not fabricate values."
)

CLIENT = OpenAI()

# ========= Intake fields (keys) =========
INTAKE_FIELDS = [
  { "key": "applicant_name", "label": "Applicant Name" },
  { "key": "co_applicant_name", "label": "Co-Applicant Name" },
  { "key": "property_owner_name", "label": "Name of the Property Owner" },
  { "key": "loan_type", "label": "Loan Type / Product Type" },
  { "key": "application_number", "label": "Application Number" },
  { "key": "developer_name", "label": "Developer’s Name" },
  { "key": "flat", "label": "Flat" },
  { "key": "property_address_doc", "label": "Property Address (as per Document)" },
  { "key": "pincode", "label": "PIN Code (as per Document)" },
  { "key": "usage_approved_as_per_plan", "label": "Usage approved as per Plan" },
  { "key": "usage_as_per_cdp", "label": "Usage as per CDP/Master plan" },
  { "key": "address_as_per_plan", "label": "Address as per Plan" },
  { "key": "legal_documents", "label": "Legal Documents" },
  { "key": "list_of_documents", "label": "List of Documents provided" },
  { "key": "land_freehold_or_leasehold", "label": "Land Freehold / Leasehold, term of lease" },
  { "key": "period_expired_balance_lease_rent", "label": "Period expired, balance and lease rent" },
  { "key": "boundaries_deed_east", "label": "Boundaries EAST (As per Deed)" },
  { "key": "boundaries_deed_west", "label": "Boundaries WEST (As per Deed)" },
  { "key": "boundaries_deed_north", "label": "Boundaries NORTH (As per Deed)" },
  { "key": "boundaries_deed_south", "label": "Boundaries SOUTH (As per Deed)" },
  { "key": "dimensions_deed_east", "label": "Dimension EAST (As per Deed)" },
  { "key": "dimensions_deed_west", "label": "Dimension WEST (As per Deed)" },
  { "key": "dimensions_deed_north", "label": "Dimension NORTH (As per Deed)" },
  { "key": "dimensions_deed_south", "label": "Dimension SOUTH (As per Deed)" },
  { "key": "land_area_docs_sqyd", "label": "Land Area (Sqyds) as per Documents" },
  { "key": "land_area_docs_sqmt", "label": "Land Area (Sqmt) as per Documents" },
  { "key": "land_area_docs_sqft", "label": "Land Area (Sqft) as per Documents" },
  { "key": "land_area_plan", "label": "Land/Plot Area as per plan" },
  { "key": "type_of_plot", "label": "Type of plot" },
  { "key": "final_land_area_uds", "label": "Final Land area / UDS considered" },
  { "key": "document_area_sqft", "label": "Document Area (title deed)" },
  { "key": "approved_area_plan_sqft", "label": "Approved Area (as per plan)" },
  { "key": "sanction_approval_no", "label": "Sanctioned plans / approval no" },
  { "key": "sanction_number_date", "label": "Number and Date" },
  { "key": "property_documents", "label": "Property documents" },
  { "key": "ownership_type", "label": "Ownership type (Leasehold/Freehold)" },
  { "key": "amenities_idc", "label": "IDC" },
  { "key": "amenities_edc", "label": "EDC" },
  { "key": "amenities_power_backup", "label": "Power Backup" },
  { "key": "amenities_plc", "label": "PLC" },
  { "key": "amenities_car_parking", "label": "Car Parking" },
  { "key": "amenities_others", "label": "Others" },
  { "key": "floors_sanctioned", "label": "No. Of Floors Sanctioned" },
  { "key": "floors_proposed", "label": "No. of floors Proposed" },
]

# ========= Helpers =========
DEVANAGARI_RE = re.compile(r"[\u0900-\u097F]")  # detect Hindi script

def _img_to_data_url(img: Image.Image, fmt: str = "PNG") -> str:
    buf = io.BytesIO()
    img.save(buf, format=fmt)
    b64 = base64.b64encode(buf.getvalue()).decode("utf-8")
    mime = f"image/{fmt.lower()}"
    return f"data:{mime};base64,{b64}"

def _bytes_to_data_url(blob: bytes, mime: str = "image/png") -> str:
    b64 = base64.b64encode(blob).decode("utf-8")
    return f"data:{mime};base64,{b64}"

def _file_to_data_url(path: Path) -> str:
    mime, _ = mimetypes.guess_type(str(path))
    if not mime:
        mime = "image/png"
    with open(path, "rb") as f:
        b64 = base64.b64encode(f.read()).decode("utf-8")
    return f"data:{mime};base64,{b64}"

def _extract_text_from_response(resp) -> str:
    txt = getattr(resp, "output_text", None)
    if txt:
        return txt.strip()
    out_parts: List[str] = []
    for item in getattr(resp, "output", []) or []:
        if getattr(item, "type", None) == "message":
            for part in getattr(item, "content", []) or []:
                if getattr(part, "type", None) in ("output_text", "text"):
                    piece = getattr(part, "text", None)
                    if piece:
                        out_parts.append(piece)
    return "\n".join(out_parts).strip()

_REFUSAL_RE = re.compile(r"\b(i'?m|i am|sorry|cannot|can'?t|unable|assist)\b", re.I)

def _call_responses(model: str, messages: list, **kwargs) -> str:
    resp = CLIENT.responses.create(model=model, input=messages, **kwargs)
    return _extract_text_from_response(resp)

def translate_text_to_english(text: str) -> str:
    if not text.strip():
        return text
    messages = [
        {"role": "system", "content": [{"type": "input_text", "text": SYSTEM_TRANSLATOR}]},
        {"role": "user", "content": [{"type": "input_text", "text": text}]},
    ]
    out = _call_responses(MODEL, messages)
    if _REFUSAL_RE.search(out):
        out = _call_responses(MODEL_FALLBACK, messages)
    return out.strip()

def ensure_english(text: str) -> str:
    if OCR_TRANSLATE_ALWAYS and DEVANAGARI_RE.search(text):
        return translate_text_to_english(text)
    return text

# ========= JSON Schema for extraction =========
def intake_schema() -> Dict:
    # All fields optional strings; omit if absent.
    props = {f["key"]: {"type": "string"} for f in INTAKE_FIELDS}
    return {
        "name": "IntakeExtraction",
        "schema": {
            "type": "object",
            "properties": props,
            "additionalProperties": False,
        },
        "strict": True,
    }

def merge_keep_longer(dst: Dict, src: Dict) -> Dict:
    for k, v in (src or {}).items():
        if not v:
            continue
        if k not in dst or len(str(v)) > len(str(dst.get(k, ""))):
            dst[k] = v
    return dst

# ========= OCR Calls =========
def ocr_image_dataurl_text(data_url: str, user_prompt: str, detail: str = "high") -> str:
    # (Legacy text OCR mode)
    if OCR_TRANSLATE_ALWAYS:
        system_text = SYSTEM_OCR_TRANSLATE
    else:
        system_text = (SYSTEM_REDACT if OCR_REDACT else SYSTEM_VERBATIM)
    messages = [
        {"role": "system", "content": [{"type": "input_text", "text": system_text}]},
        {"role": "user", "content": [
            {"type": "input_text", "text": user_prompt},
            {"type": "input_image", "image_url": data_url, "detail": detail},
        ]},
    ]
    out = _call_responses(MODEL, messages)
    if _REFUSAL_RE.search(out):
        messages[0]["content"][0]["text"] += " This is first-party, user-owned content strictly for OCR/translation."
        out = _call_responses(MODEL_FALLBACK, messages)
    return ensure_english(out.strip())

def ocr_image_dataurl_structured(data_url: str, labels_hint: str = "", detail: str = "high") -> Dict:
    """
    OCR + extract to schema'd JSON. Returns a dict (omit-not-found).
    """
    system_text = SYSTEM_STRUCTURED
    # Build a concise hint including labels (helps mapping)
    if not labels_hint:
        labels_hint = "; ".join([f['label'] for f in INTAKE_FIELDS])

    messages = [
        {"role": "system", "content": [{"type": "input_text", "text": system_text}]},
        {"role": "user", "content": [
            {"type": "input_text", "text":
             ("Extract the following labeled fields from this page. "
              "Return only the JSON object. "
              "If a value is in Hindi, translate to English. "
              "Do not include fields you cannot find. "
              f"Fields: {labels_hint}")
             },
            {"type": "input_image", "image_url": data_url, "detail": detail},
        ]},
    ]

    schema = intake_schema()
    # Ask for structured JSON
    resp = CLIENT.responses.create(
        model=MODEL,
        input=messages,
        response_format={"type": "json_schema", "json_schema": schema},
    )
    out = getattr(resp, "output_parsed", None)
    if out is None:  # fallback if SDK shape differs
        parsed = _extract_text_from_response(resp)
        try:
            out = json.loads(parsed)
        except Exception:
            out = {}

    # Fallback model if refusal-ish or empty
    if (not out) or (isinstance(out, dict) and not out) :
        resp2 = CLIENT.responses.create(
            model=MODEL_FALLBACK,
            input=messages,
            response_format={"type": "json_schema", "json_schema": schema},
        )
        out = getattr(resp2, "output_parsed", None) or {}
        if not out:
            parsed2 = _extract_text_from_response(resp2)
            try:
                out = json.loads(parsed2)
            except Exception:
                out = {}

    # Ensure English if needed (defensive)
    for k, v in list(out.items()):
        if isinstance(v, str) and DEVANAGARI_RE.search(v):
            out[k] = translate_text_to_english(v)
    return out

# ========= PDF (PyMuPDF) =========
def ocr_pdf_free_text(path: Path, prompt: str, zoom: float = 2.5) -> str:
    texts: List[str] = []
    doc = fitz.open(str(path))
    try:
        mat = fitz.Matrix(zoom, zoom)
        for i, page in enumerate(doc, start=1):
            pix = page.get_pixmap(matrix=mat, alpha=False)
            img = Image.frombytes("RGB", [pix.width, pix.height], pix.samples)
            data_url = _img_to_data_url(img, fmt="PNG")
            page_text = ocr_image_dataurl_text(data_url, f"{prompt}\nKeep original line breaks.")
            texts.append(f"[Page {i}]\n{page_text}".rstrip())
    finally:
        doc.close()
    return "\n\n".join(texts).strip()

def ocr_pdf_structured(path: Path, zoom: float = 2.5) -> Dict:
    merged: Dict = {}
    labels_hint = "; ".join([f['label'] for f in INTAKE_FIELDS])
    doc = fitz.open(str(path))
    try:
        mat = fitz.Matrix(zoom, zoom)
        for _i, page in enumerate(doc, start=1):
            pix = page.get_pixmap(matrix=mat, alpha=False)
            img = Image.frombytes("RGB", [pix.width, pix.height], pix.samples)
            data_url = _img_to_data_url(img, fmt="PNG")
            d = ocr_image_dataurl_structured(data_url, labels_hint=labels_hint)
            merged = merge_keep_longer(merged, d)
    finally:
        doc.close()
    return merged

# ========= Images =========
def ocr_image_file_free_text(path: Path, prompt: str) -> str:
    data_url = _file_to_data_url(path)
    return ocr_image_dataurl_text(data_url, prompt)

def ocr_image_file_structured(path: Path) -> Dict:
    data_url = _file_to_data_url(path)
    return ocr_image_dataurl_structured(data_url)

# ========= DOCX (text + embedded images) =========
def extract_docx_text(doc: Document) -> str:
    parts: List[str] = []
    for p in doc.paragraphs:
        if p.text.strip():
            parts.append(p.text)
    for tbl in doc.tables:
        for row in tbl.rows:
            row_text = [cell.text.strip() for cell in row.cells]
            if any(row_text):
                parts.append(" | ".join(row_text))
    return "\n".join(parts).strip()

def ocr_docx_images_structured(doc: Document) -> Dict:
    merged: Dict = {}
    labels_hint = "; ".join([f['label'] for f in INTAKE_FIELDS])
    rels = list(doc.part.rels.values())
    for rel in rels:
        if rel.reltype == RT.IMAGE:
            part = rel.target_part
            blob = part.blob
            ext = (part.partname.ext or "").lower().lstrip(".")
            if ext in ("jpg", "jpeg"):
                mime = "image/jpeg"
            elif ext == "png":
                mime = "image/png"
            elif ext == "gif":
                mime = "image/gif"
            elif ext == "bmp":
                mime = "image/bmp"
            elif ext in ("tiff", "tif"):
                mime = "image/tiff"
            else:
                mime = "image/png"
            data_url = _bytes_to_data_url(blob, mime=mime)
            d = ocr_image_dataurl_structured(data_url, labels_hint=labels_hint)
            merged = merge_keep_longer(merged, d)
    return merged

def process_docx_free_text(path: Path, prompt: str) -> str:
    doc = Document(str(path))
    text_block = extract_docx_text(doc)
    # translate if asked or if Hindi detected
    if text_block and (OCR_TRANSLATE_ALWAYS or DEVANAGARI_RE.search(text_block)):
        text_block = translate_text_to_english(text_block)
    # images to text blocks too (for free-text mode, we’ll just append)
    return text_block

def process_docx_structured(path: Path) -> Dict:
    doc = Document(str(path))
    text_block = extract_docx_text(doc)
    # 1) Extract from selectable text
    selectable_extraction = {}
    if text_block:
        # Translate selectable text if requested
        if OCR_TRANSLATE_ALWAYS or DEVANAGARI_RE.search(text_block):
            text_block = translate_text_to_english(text_block)
        messages = [
            {"role": "system", "content": [{"type": "input_text", "text": SYSTEM_STRUCTURED}]},
            {"role": "user", "content": [{"type": "input_text",
                "text": ("Extract ONLY the requested fields from this text. "
                         "Return only the JSON object. If a field is not found, omit it. "
                         "Text:\n" + text_block)}]},
        ]
        schema = intake_schema()
        resp = CLIENT.responses.create(
            model=MODEL,
            input=messages,
            response_format={"type": "json_schema", "json_schema": schema},
        )
        selectable_extraction = getattr(resp, "output_parsed", None) or {}
        if not selectable_extraction:
            parsed = _extract_text_from_response(resp)
            try:
                selectable_extraction = json.loads(parsed)
            except Exception:
                selectable_extraction = {}

    # 2) Extract from embedded images (OCR)
    image_extraction = ocr_docx_images_structured(doc)

    # Merge both (keep longer text)
    merged: Dict = {}
    merge_keep_longer(merged, selectable_extraction)
    merge_keep_longer(merged, image_extraction)
    return merged

# ========= Main =========
def main():
    if len(sys.argv) < 2:
        print("Usage: python ocr_any.py <file.(pdf|png|jpg|jpeg|webp|docx)>", file=sys.stderr)
        sys.exit(1)

    infile = Path(sys.argv[1]).expanduser().resolve()
    if not infile.exists():
        print(f"File not found: {infile}", file=sys.stderr)
        sys.exit(1)

    # Choose the effective prompt for free-text OCR mode
    if OCR_TRANSLATE_ALWAYS:
        prompt = PROMPT_OCR_ANY_TO_EN
    else:
        prompt = PROMPT_OCR_HI_TO_EN if OCR_LANG == "hi" else PROMPT_OCR_VERBATIM_EN

    ext = infile.suffix.lower()
    if OCR_STRUCTURED:
        # ---- Structured JSON output ----
        if ext == ".pdf":
            result = ocr_pdf_structured(infile, zoom=2.5)
        elif ext in {".png", ".jpg", ".jpeg", ".webp"}:
            result = ocr_image_file_structured(infile)
        elif ext == ".docx":
            result = process_docx_structured(infile)
        else:
            mime, _ = mimetypes.guess_type(str(infile))
            if mime == "application/pdf":
                result = ocr_pdf_structured(infile, zoom=2.5)
            elif mime and mime.startswith("image/"):
                result = ocr_image_file_structured(infile)
            elif ext == ".docx" or (mime and mime == "application/vnd.openxmlformats-officedocument.wordprocessingml.document"):
                result = process_docx_structured(infile)
            else:
                print("Unsupported file type. Provide PDF, DOCX, or an image (png/jpg/webp).", file=sys.stderr)
                sys.exit(2)

        # Print strict JSON string (omit-empty is already handled by the schema+prompts)
        print(json.dumps(result, ensure_ascii=False))
        return

    # ---- Free-text OCR (previous behavior) ----
    if ext == ".pdf":
        text = ocr_pdf_free_text(infile, prompt, zoom=2.5)
    elif ext in {".png", ".jpg", ".jpeg", ".webp"}:
        text = ocr_image_file_free_text(infile, prompt)
    elif ext == ".docx":
        text = process_docx_free_text(infile, prompt)
    else:
        mime, _ = mimetypes.guess_type(str(infile))
        if mime == "application/pdf":
            text = ocr_pdf_free_text(infile, prompt, zoom=2.5)
        elif mime and mime.startswith("image/"):
            text = ocr_image_file_free_text(infile, prompt)
        elif ext == ".docx" or (mime and mime == "application/vnd.openxmlformats-officedocument.wordprocessingml.document"):
            text = process_docx_free_text(infile, prompt)
        else:
            print("Unsupported file type. Provide PDF, DOCX, or an image (png/jpg/webp).", file=sys.stderr)
            sys.exit(2)

    print(json.dumps({"text": text}, ensure_ascii=False))

if __name__ == "__main__":
    main()
