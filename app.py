# app.py
import io
import os
import re
import json
import time
import base64
import logging
import mimetypes
import uuid
import contextvars
from pathlib import Path
from typing import Dict, List, Tuple, Optional

import fitz  # PyMuPDF
from PIL import Image
from fastapi import FastAPI, UploadFile, File, Form, HTTPException, Depends, status, Request
from fastapi.responses import PlainTextResponse
from fastapi.security import APIKeyHeader, HTTPBearer, HTTPAuthorizationCredentials
from fastapi.middleware.cors import CORSMiddleware
from docx import Document
from docx.opc.constants import RELATIONSHIP_TYPE as RT
from openai import OpenAI

# ===== NEW: Google Vision / Storage =====
from google.cloud import vision
from google.cloud.vision_v1.types import (
    Feature,
    InputConfig,
    OutputConfig,
    GcsSource,
    GcsDestination,
    AsyncAnnotateFileRequest,
)
from google.cloud import storage
from google.api_core.operation import Operation

# =========================
# Logging setup
# =========================
LOG_LEVEL = os.getenv("LOG_LEVEL", "INFO").upper()
logging.basicConfig(
    level=getattr(logging, LOG_LEVEL, logging.INFO),
    format="%(asctime)s %(levelname)s %(name)s - %(message)s",
)
logger = logging.getLogger("ocr_engine")

# Request correlation id (per-request)
REQ_ID: contextvars.ContextVar[str] = contextvars.ContextVar("req_id", default="-")
def rid() -> str:
    return REQ_ID.get()

def log(level: int, msg: str, **kv):
    kv_str = " ".join(f"{k}={v}" for k, v in kv.items())
    logger.log(level, f"[rid={rid()}] {msg}" + (f" | {kv_str}" if kv_str else ""))

app = FastAPI(title="OCR Intake Parser", version="2.1.0-gcv")

# ---- CORS ----
app.add_middleware(
    CORSMiddleware,
    allow_origins=[
        "http://localhost:5173",
        "http://127.0.0.1:5173",
        # add prod origins/domains here as needed
        "https://ocr.finotoai.com",
    ],
    allow_credentials=False,
    allow_methods=["POST", "GET", "OPTIONS"],
    allow_headers=[
        "Authorization",
        "X-API-Key",
        "Content-Type",
        "Accept",
        # optionally: "X-Requested-With"
    ],
    expose_headers=[],
    max_age=86400,
)

# =========================
# Config / Model selection
# =========================
# OpenAI: text-only model for interpretation stays the same
MODEL = os.getenv("OPENAI_MODEL", "gpt-5")  # kept for fallback only
TEXT_MODEL = os.getenv("OPENAI_TEXT_MODEL") or os.getenv("OPENAI_MODEL_FALLBACK", "gpt-4.1-mini")
CLIENT = OpenAI(timeout=600.0, max_retries=5)
log(logging.INFO, "Server startup", openai_timeout=600, openai_retries=5, text_model=TEXT_MODEL)

# Google Vision defaults
GCV_BUCKET = os.getenv("GCV_BUCKET", "ocr-storage-db").strip()
GCV_PREFIX = os.getenv("GCV_PREFIX", "uploads/").strip().strip("/")  # e.g. "uploads"
if GCV_PREFIX:
    GCV_PREFIX = GCV_PREFIX + "/"

# Init Vision + Storage clients (will raise if credentials missing)
try:
    VISION_CLIENT = vision.ImageAnnotatorClient()
    STORAGE_CLIENT = storage.Client()
    log(logging.INFO, "Google clients initialized", gcv_bucket=GCV_BUCKET, gcv_prefix=GCV_PREFIX or "(none)")
except Exception as e:
    VISION_CLIENT = None
    STORAGE_CLIENT = None
    log(logging.ERROR, "Google clients init failed", error=str(e))

# =========================
# Auth (SECRET_API_KEY)
# =========================
SECRET_API_KEY = os.getenv("SECRET_API_KEY")
_api_key_header = APIKeyHeader(name="X-API-Key", auto_error=False)
_http_bearer = HTTPBearer(auto_error=False)

def require_api_key(
    api_key_header: Optional[str] = Depends(_api_key_header),
    bearer: Optional[HTTPAuthorizationCredentials] = Depends(_http_bearer),
):
    supplied = None
    if api_key_header:
        supplied = api_key_header.strip()
    elif bearer and bearer.scheme.lower() == "bearer":
        supplied = (bearer.credentials or "").strip()
    if not SECRET_API_KEY:
        log(logging.ERROR, "Auth failed: SECRET_API_KEY not configured")
        raise HTTPException(
            status_code=status.HTTP_503_SERVICE_UNAVAILABLE,
            detail="Server not configured with SECRET_API_KEY",
        )
    if not supplied or supplied != SECRET_API_KEY:
        log(logging.WARNING, "Auth failed: invalid or missing API key")
        raise HTTPException(
            status_code=status.HTTP_401_UNAUTHORIZED,
            detail="Invalid or missing API key",
        )
    return True

# =========================
# Helpers
# =========================
NON_ALNUM = re.compile(r"[^a-z0-9]+")
DEVANAGARI_RE = re.compile(r"[\u0900-\u097F]")
JSON_OBJ_RE = re.compile(r"\{.*\}", re.DOTALL)

def label_to_key(label: str) -> str:
    s = (label or "").strip().lower()
    s = NON_ALNUM.sub("_", s)
    s = re.sub(r"_+", "_", s).strip("_")
    return (s or "field")[:80]

def intake_schema_rich(fields: List[str]) -> Dict:
    entries = [{"key": label_to_key(lbl), "label": lbl} for lbl in fields]
    field_obj = {
        "type": "object",
        "properties": {
            "value":  {"type": "string"},
            "conf":   {"type": "number", "minimum": 0.0, "maximum": 1.0},
            "source": {"type": "string"},
        },
        "required": ["value", "conf", "source"],
        "additionalProperties": False,
    }
    props = {e["key"]: field_obj for e in entries}
    schema = {
        "name": "IntakeExtractionWithConfidence",
        "schema": {"type": "object", "properties": props, "additionalProperties": False},
        "strict": True,
    }
    log(logging.DEBUG, "Built JSON schema", fields=len(fields), keys=list(props.keys()))
    return schema

def first_json_object(text: str) -> Dict:
    if not text:
        return {}
    m = JSON_OBJ_RE.search(text)
    if not m:
        return {}
    try:
        return json.loads(m.group(0))
    except Exception as e:
        log(logging.WARNING, "JSON parse (regex) failed", error=str(e))
        return {}

def clamp_conf(x):
    try:
        v = float(x)
    except Exception:
        return 0.5
    return 0.0 if v < 0 else 1.0 if v > 1 else v

def extract_text_from_response(resp) -> str:
    txt = getattr(resp, "output_text", None)
    if txt:
        return txt.strip()
    out_parts: List[str] = []
    for item in getattr(resp, "output", []) or []:
        if getattr(item, "type", None) == "message":
            for part in getattr(item, "content", []) or []:
                if getattr(part, "type", None) in ("output_text", "text"):
                    piece = getattr(part, "text", None)
                    if piece:
                        out_parts.append(piece)
    return "\n".join(out_parts).strip()

def call_responses(messages: list, *, model: str) -> str:
    t0 = time.time()
    try:
        log(logging.INFO, "OpenAI call:start", model=model, stage="responses.create", msg_roles=[m["role"] for m in messages])
        resp = CLIENT.responses.create(model=model, input=messages)
        out = extract_text_from_response(resp)
        log(logging.INFO, "OpenAI call:done", model=model, ms=int((time.time()-t0)*1000), out_len=len(out))
        return out
    except Exception as e:
        log(logging.ERROR, "OpenAI call:error", model=model, ms=int((time.time()-t0)*1000), error=str(e))
        raise

def translate_to_english(text: str) -> str:
    if not text.strip():
        return text
    messages = [
        {"role": "system", "content": [{"type": "input_text", "text":
            "You are a professional translator. Translate to natural, fluent English. Preserve line breaks; no commentary."}]},
        {"role": "user", "content": [{"type": "input_text", "text": text}]},
    ]
    log(logging.INFO, "Translate:start", model=TEXT_MODEL, detected_script="devanagari")
    out = call_responses(messages, model=TEXT_MODEL)
    if not out or out.lower().startswith(("sorry", "i can't", "i cannot", "i am sorry")):
        log(logging.WARNING, "Translate:fallback_to_vision_model", fallback_model=MODEL)
        out = call_responses(messages, model=MODEL)  # fallback
    log(logging.INFO, "Translate:done", out_len=len(out))
    return out.strip()

def text_maybe_translate(text: str, translate: bool) -> str:
    if translate and DEVANAGARI_RE.search(text or ""):
        return translate_to_english(text)
    return text

# =========================
# Google Vision helpers (from google_vision.py, adapted)
# =========================
# Break types
_BREAK_SPACE = 1
_BREAK_EOL_SURE = 2
_BREAK_SURE_SPACE = 3
_BREAK_LINE_BREAK = 5

def _word_text_proto(word) -> str:
    return "".join(s.text or "" for s in word.symbols)

def _word_break_proto(word) -> int | None:
    if not word.symbols:
        return None
    last = word.symbols[-1]
    if last.property and last.property.detected_break:
        return last.property.detected_break.type
    return None

def _paragraph_lines_proto(paragraph) -> List[str]:
    lines: List[str] = []
    buf: List[str] = []
    for w in paragraph.words:
        wtxt = _word_text_proto(w)
        if not wtxt:
            continue
        if buf and not buf[-1].endswith(" "):
            buf.append(" ")
        buf.append(wtxt)
        br = _word_break_proto(w)
        if br in (_BREAK_SPACE, _BREAK_SURE_SPACE):
            if not buf[-1].endswith(" "):
                buf.append(" ")
        elif br in (_BREAK_EOL_SURE, _BREAK_LINE_BREAK):
            line = "".join(buf).rstrip()
            if line:
                lines.append(line)
            buf = []
    tail = "".join(buf).strip()
    if tail:
        lines.append(tail)
    return lines

def _pages_to_map_proto(pages) -> Tuple[str, int, Dict[int, List[str]]]:
    per_page: Dict[int, List[str]] = {}
    all_lines: List[str] = []
    for idx, page in enumerate(pages, start=1):
        lines: List[str] = []
        for block in page.blocks:
            for para in block.paragraphs:
                lines.extend(_paragraph_lines_proto(para))
        per_page[idx] = lines
        all_lines.extend(lines)
    joined = "\n".join(all_lines)
    return joined, len(per_page), per_page

def _gcs_upload(storage_client: storage.Client, bucket: str, key: str, data: bytes, content_type: str):
    b = storage_client.bucket(bucket)
    b.storage_class  # touch to surface permission errors early
    blob = b.blob(key)
    blob.content_type = content_type
    blob.upload_from_string(data, content_type=content_type)

def _gcs_list_prefix(storage_client: storage.Client, bucket: str, prefix: str):
    return list(storage_client.list_blobs(bucket, prefix=prefix))

def _gcs_download_text(storage_client: storage.Client, bucket: str, blob_name: str) -> str:
    b = storage_client.bucket(bucket)
    return b.blob(blob_name).download_as_text()

def _vision_pdf_async(client: vision.ImageAnnotatorClient,
                      storage_client: storage.Client,
                      bucket: str,
                      gcs_input_uri: str,
                      gcs_output_prefix: str,
                      timeout: int = 60*30):
    feature = Feature(type_=Feature.Type.DOCUMENT_TEXT_DETECTION)
    input_cfg = InputConfig(gcs_source=GcsSource(uri=gcs_input_uri), mime_type="application/pdf")
    output_cfg = OutputConfig(gcs_destination=GcsDestination(uri=gcs_output_prefix), batch_size=25)
    request = AsyncAnnotateFileRequest(features=[feature], input_config=input_cfg, output_config=output_cfg)
    op: Operation = client.async_batch_annotate_files(requests=[request])
    op.result(timeout=timeout)  # wait for completion

    out_bucket = gcs_output_prefix.split("gs://", 1)[1].split("/", 1)[0]
    out_prefix = gcs_output_prefix.split(out_bucket, 1)[1].lstrip("/")
    blobs = _gcs_list_prefix(storage_client, out_bucket, out_prefix)
    blobs = [b for b in blobs if b.name.endswith(".json")]
    blobs.sort(key=lambda b: b.name)

    # Helpers for JSON (dict) structure
    def _word_text_dict(w: dict) -> str:
        return "".join(s.get("text", "") for s in (w.get("symbols") or []))

    def _word_break_dict(w: dict):
        syms = w.get("symbols") or []
        if not syms:
            return None
        last = syms[-1]
        prop = last.get("property") or {}
        return (prop.get("detectedBreak") or {}).get("type")

    per_page: Dict[int, List[str]] = {}
    all_lines: List[str] = []
    page_idx = 0

    for blob in blobs:
        data = json.loads(_gcs_download_text(storage_client, out_bucket, blob.name))
        for r in data.get("responses", []):
            ann = r.get("fullTextAnnotation") or {}
            pages = ann.get("pages") or []
            for p in pages:
                page_idx += 1
                lines: List[str] = []
                buf: List[str] = []
                for block in p.get("blocks", []):
                    for para in block.get("paragraphs", []):
                        for w in para.get("words", []):
                            wtxt = _word_text_dict(w)
                            if not wtxt:
                                continue
                            if buf and not buf[-1].endswith(" "):
                                buf.append(" ")
                            buf.append(wtxt)
                            br = _word_break_dict(w)
                            if br in (_BREAK_SPACE, _BREAK_SURE_SPACE):
                                if not buf[-1].endswith(" "):
                                    buf.append(" ")
                            elif br in (_BREAK_EOL_SURE, _BREAK_LINE_BREAK):
                                line = "".join(buf).rstrip()
                                if line:
                                    lines.append(line)
                                buf = []
                        if buf:
                            line = "".join(buf).rstrip()
                            if line:
                                lines.append(line)
                            buf = []
                per_page[page_idx] = lines
                all_lines.extend(lines)

    joined = "\n".join(all_lines)
    page_count = page_idx or 1
    return joined, page_count, per_page

# =========================
# STAGE 1: OCR (Vision-first)
# =========================
def ocr_pdf_pages(blob: bytes, lang: str) -> List[Tuple[int, str]]:
    """
    Default: Google Cloud Vision (async PDF) via GCS (per google_vision.py).
    Returns per-page text list [(page_number, text), ...].
    """
    if not VISION_CLIENT or not STORAGE_CLIENT:
        raise HTTPException(status_code=503, detail="Google Vision not initialized (check GOOGLE_APPLICATION_CREDENTIALS).")
    if not GCV_BUCKET:
        raise HTTPException(status_code=503, detail="GCV_BUCKET env is required for PDF OCR.")

    translate = (lang or "en").lower() == "hi"
    key = f"{GCV_PREFIX}{uuid.uuid4().hex}.pdf"
    _gcs_upload(STORAGE_CLIENT, GCV_BUCKET, key, blob, "application/pdf")
    gcs_in = f"gs://{GCV_BUCKET}/{key}"
    out_prefix = f"gs://{GCV_BUCKET}/{GCV_PREFIX}vision_out/{uuid.uuid4().hex}/"

    log(logging.INFO, "GCV PDF OCR: start", input=gcs_in, out_prefix=out_prefix)
    joined, page_count, per_page_map = _vision_pdf_async(
        VISION_CLIENT, STORAGE_CLIENT, GCV_BUCKET, gcs_in, out_prefix, timeout=60*30
    )
    log(logging.INFO, "GCV PDF OCR: done", pages=page_count, chars=len(joined))

    results: List[Tuple[int, str]] = []
    for pno in range(1, page_count + 1):
        txt = "\n".join(per_page_map.get(pno, []))
        txt = text_maybe_translate(txt, translate)
        results.append((pno, txt))
    return results

def ocr_image_pages(filename: str, blob: bytes, lang: str) -> List[Tuple[int, str]]:
    """
    Default: Google Cloud Vision document_text_detection for single image.
    """
    if not VISION_CLIENT:
        raise HTTPException(status_code=503, detail="Google Vision not initialized (check GOOGLE_APPLICATION_CREDENTIALS).")

    translate = (lang or "en").lower() == "hi"
    image = vision.Image(content=blob)
    log(logging.INFO, "GCV Image OCR: start", filename=filename, bytes=len(blob))
    resp = VISION_CLIENT.document_text_detection(
        image=image, image_context={"language_hints": ["hi", "en"]}
    )
    if resp.error.message:
        raise HTTPException(status_code=502, detail=f"Vision error: {resp.error.message}")

    ann = resp.full_text_annotation
    if not ann or not ann.pages:
        pages = [(1, "")]
    else:
        _, page_count, per_page_map = _pages_to_map_proto(ann.pages)
        pages = [(pno, "\n".join(per_page_map.get(pno, []))) for pno in range(1, page_count + 1)]

    pages = [(pno, text_maybe_translate(txt, translate)) for (pno, txt) in pages]
    log(logging.INFO, "GCV Image OCR: done", pages=len(pages))
    return pages

def ocr_docx_pages(blob: bytes, lang: str) -> List[Tuple[int, str]]:
    translate = (lang or "en").lower() == "hi"
    log(logging.INFO, "DOCX parse:start")
    doc = Document(io.BytesIO(blob))
    parts: List[str] = []
    for p in doc.paragraphs:
        if p.text.strip():
            parts.append(p.text)
    for tbl in doc.tables:
        for row in tbl.rows:
            row_text = [cell.text.strip() for cell in row.cells]
            if any(row_text):
                parts.append(" | ".join(row_text))
    text = "\n".join(parts).strip()
    log(logging.INFO, "DOCX parse:done", chars=len(text))
    text = text_maybe_translate(text, translate)
    return [(1, text)]

def ocr_pages(filename: str, blob: bytes, lang: str) -> List[Tuple[int, str]]:
    """
    Vision-first dispatch. DOCX stays local. OpenAI is NOT used for OCR anymore.
    """
    ext = Path(filename).suffix.lower()
    if ext == ".pdf" or blob[:5] == b"%PDF-":
        log(logging.INFO, "OCR stage: PDF detected (Google Vision)")
        return ocr_pdf_pages(blob, lang)
    if ext == ".docx" or mimetypes.guess_type(filename)[0] == \
        "application/vnd.openxmlformats-officedocument.wordprocessingml.document":
        log(logging.INFO, "OCR stage: DOCX detected (local text extraction)")
        return ocr_docx_pages(blob, lang)
    mime, _ = mimetypes.guess_type(filename)
    if mime and mime.startswith("image/"):
        log(logging.INFO, "OCR stage: IMAGE detected (Google Vision)", mime=mime)
        return ocr_image_pages(filename, blob, lang)
    log(logging.INFO, "OCR stage: defaulting to IMAGE path (Google Vision)")
    return ocr_image_pages(filename, blob, lang)

# =========================
# STAGE 2: INTERPRETATION (text-only, unchanged)
# =========================
DOC_INSTR_BASE = (
    "You are an information extraction engine. Use ONLY the provided page texts. "
    "For each requested field, return an object with keys: value (string), conf (0-1), "
    "and source (format 'FILENAME#page N'). Omit fields not confidently present. Do not fabricate."
)

def build_interpret_messages(filename: str, pages: List[Tuple[int, str]], fields: List[str]) -> List[Dict]:
    entries = [{"key": label_to_key(lbl), "label": lbl} for lbl in fields]
    labels_hint = "; ".join([e["label"] for e in entries])
    mapping_lines = "\n".join([f"- {e['label']} -> {e['key']}" for e in entries])

    blocks = []
    for pno, txt in pages:
        if txt and txt.strip():
            blocks.append(f"Page {pno}:\n{txt.strip()}")

    user_intro = (
        f"Document name: {filename}.\n"
        f"Extract only these fields (labels): {labels_hint}\n"
        "Use the following key names in JSON (label -> key):\n"
        f"{mapping_lines}\n"
        "Always include 'source' as 'FILENAME#page N' where N is the page number from which you derived the value."
    )
    user_text = user_intro + "\n\n==== Page Texts ====\n" + ("\n\n---\n\n".join(blocks) if blocks else "Page 1:\n")

    log(logging.INFO, "Interpretation:build_messages", pages=len(pages), fields=len(fields))
    return [
        {"role": "system", "content": [{"type": "input_text", "text": DOC_INSTR_BASE}]},
        {"role": "user", "content": [{"type": "input_text", "text": user_text}]},
    ]

def structured_from_messages(messages: list, schema: dict, *, model: str) -> Dict:
    # Prefer JSON schema; fall back to strict JSON instruction
    t0 = time.time()
    try:
        log(logging.INFO, "Interpretation:OpenAI JSON-schema:start", model=model)
        resp = CLIENT.responses.create(
            model=model,
            input=messages,
            response_format={"type": "json_schema", "json_schema": schema},
        )
        out = getattr(resp, "output_parsed", None)
        if out is not None:
            log(logging.INFO, "Interpretation:OpenAI JSON-schema:done", model=model, ms=int((time.time()-t0)*1000), keys=list(out.keys()))
            return out
        raw = extract_text_from_response(resp)
        out2 = first_json_object(raw)
        log(logging.INFO, "Interpretation:OpenAI JSON-schema:parsed_text", model=model, ms=int((time.time()-t0)*1000), parsed=bool(out2))
        return out2
    except TypeError:
        log(logging.WARNING, "Interpretation:JSON-schema not supported by SDK, falling back", model=model)
        strong = list(messages)
        for blk in strong[-1]["content"]:
            if blk.get("type") == "input_text":
                blk["text"] += (
                    "\n\nReturn ONLY a minified JSON object (no prose, no code fences). "
                    "For each present field include: {\"value\": string, \"conf\": number 0-1, \"source\": string}."
                )
                break
        raw = call_responses(strong, model=model)
        out = first_json_object(raw)
        log(logging.INFO, "Interpretation:fallback parsed", model=model, parsed=bool(out))
        if out:
            return out
        # final fallback to vision model if text model fails
        log(logging.WARNING, "Interpretation:fallback to vision model", model=MODEL)
        raw2 = call_responses(strong, model=MODEL)
        return first_json_object(raw2) or {}

def interpret_text_structured(filename: str, pages: List[Tuple[int, str]], fields: List[str]) -> Dict:
    if not fields:
        log(logging.WARNING, "Interpretation: no fields provided")
        return {}
    schema = intake_schema_rich(fields)
    messages = build_interpret_messages(filename, pages, fields)
    out = structured_from_messages(messages, schema, model=TEXT_MODEL)

    # Normalize conf/source
    for k, obj in list(out.items()):
        if not isinstance(obj, dict):
            out[k] = {"value": str(obj), "conf": 0.5, "source": f"{filename}#page 1"}
            obj = out[k]
        obj["conf"] = clamp_conf(obj.get("conf", 0.5))
        if not isinstance(obj.get("source"), str) or "#page" not in obj.get("source", ""):
            guess = 1
            val = (obj.get("value") or "").strip()
            if val:
                for pno, txt in pages:
                    if val and (val in (txt or "")):
                        guess = pno
                        break
            obj["source"] = f"{filename}#page {guess}"
    log(logging.INFO, "Interpretation:done", fields_found=len(out))
    return out

# =========================
# Pipeline
# =========================
def detect_type(filename: str, blob: bytes) -> str:
    ext = Path(filename).suffix.lower()
    if ext == ".pdf" or blob[:5] == b"%PDF-":
        return "pdf"
    if ext == ".docx" or mimetypes.guess_type(filename)[0] == \
        "application/vnd.openxmlformats-officedocument.wordprocessingml.document":
        return "docx"
    mime, _ = mimetypes.guess_type(filename)
    if mime and mime.startswith("image/"):
        return "image"
    return "image"

def parse_file(filename: str, blob: bytes, lang: str, fields: List[str]) -> Dict:
    t0 = time.time()
    ftype = detect_type(filename, blob)
    log(logging.INFO, "Parse file:start", filename=filename, size=len(blob), type=ftype, lang=lang, fields=len(fields))
    # 1) OCR pages → texts (NOW Vision-first)
    pages = ocr_pages(filename, blob, lang)
    log(logging.INFO, "Parse file:OCR done", filename=filename, pages=len(pages), ms=int((time.time()-t0)*1000))
    # 2) Interpretation on text only (unchanged)
    t1 = time.time()
    result = interpret_text_structured(filename, pages, fields)
    log(logging.INFO, "Parse file:Interpretation done", filename=filename, ms=int((time.time()-t1)*1000))
    return result

# =========================
# API: /parse
# =========================
@app.post("/parse", response_class=PlainTextResponse)
async def parse(
    request: Request,
    files: List[UploadFile] = File(..., description="List of files (PDF/PNG/JPG/WEBP/DOCX)"),
    langs: List[str] = Form(..., description="List of language tags (en/hi) in the same order as files"),
    fields_to_extract: List[str] = Form(
        ...,
        description="Mandatory: field labels to extract (repeat this field OR provide a comma/newline/semicolon-separated list)",
    ),
    _auth_ok: bool = Depends(require_api_key),
) -> PlainTextResponse:
    # Correlate logs for this request
    REQ_ID.set(str(uuid.uuid4())[:8])
    log(logging.INFO, "Request:start", client=str(request.client))

    # Normalize files
    files = [f for f in files if getattr(f, "filename", None)]
    log(logging.INFO, "Request:files received", count=len(files), names=[f.filename for f in files])

    # Normalize langs (supports 'hi,en' or repeated fields)
    if len(langs) == 1:
        langs = [s.strip() for s in re.split(r"[,\s]+", langs[0]) if s.strip()]
    langs = [l.strip().lower() for l in langs if l and l.strip()]
    if len(langs) < len(files):
        langs += ["en"] * (len(files) - len(langs))
    elif len(langs) > len(files):
        langs = langs[:len(files)]
    bad = [l for l in langs if l not in {"en", "hi"}]
    if bad:
        log(logging.WARNING, "Request:bad langs", langs=langs)
        raise HTTPException(status_code=400, detail=f"Invalid language(s): {bad}. Use 'en' or 'hi'.")
    log(logging.INFO, "Request:langs normalized", langs=langs)

    # Mandatory fields_to_extract
    if not fields_to_extract:
        log(logging.WARNING, "Request:missing fields_to_extract")
        raise HTTPException(status_code=400, detail="fields_to_extract is required and cannot be empty.")
    if len(fields_to_extract) == 1:
        labels = [s.strip() for s in re.split(r"[,\n;]+", fields_to_extract[0]) if s.strip()]
    else:
        labels = []
        for item in fields_to_extract:
            labels.extend([s.strip() for s in re.split(r"[,\n;]+", item) if s.strip()])
    fields = [lbl for lbl in labels if lbl]
    if not fields:
        log(logging.WARNING, "Request:empty fields_to_extract after parsing")
        raise HTTPException(status_code=400, detail="fields_to_extract must contain at least one non-empty label.")
    log(logging.INFO, "Request:fields_to_extract parsed", count=len(fields))

    # Process files
    result: Dict[str, str] = {}
    for idx, uf in enumerate(files):
        t_file = time.time()
        blob = await uf.read()
        log(logging.INFO, "File read", filename=uf.filename, bytes=len(blob))
        lang = langs[idx]
        try:
            parsed_obj = parse_file(uf.filename, blob, lang, fields)
        except Exception as e:
            log(logging.ERROR, "Parse error", filename=uf.filename, error=str(e))
            raise
        # Contract: filename -> JSON string (minified) of extracted fields
        js = json.dumps(parsed_obj, ensure_ascii=False)
        result[uf.filename] = js
        log(logging.INFO, "File done", filename=uf.filename, out_chars=len(js), ms=int((time.time()-t_file)*1000))

    log(logging.INFO, "Request:done", files=len(files))
    return PlainTextResponse(content=json.dumps(result, ensure_ascii=False), media_type="application/json")

# Optional health for liveness checks
@app.get("/health")
def health():
    ok = bool(SECRET_API_KEY)
    gcv_ok = bool(VISION_CLIENT) and bool(STORAGE_CLIENT)
    return {
        "ok": ok and gcv_ok,
        "auth": bool(SECRET_API_KEY),
        "gcv_clients": gcv_ok,
        "gcv_bucket": bool(GCV_BUCKET),
        "text_model": TEXT_MODEL,
    }
